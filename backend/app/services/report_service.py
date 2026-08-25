"""
Report generation service — produces PDF (via ReportLab), DOCX, Markdown, and HTML exports.
Includes on-demand generation fallback so downloads never fail even for historical reports.
"""
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional
from uuid import UUID

import structlog
from app.core.config import settings

logger = structlog.get_logger(__name__)

REPORTS_DIR = Path(settings.REPORTS_DIR)
REPORTS_DIR.mkdir(parents=True, exist_ok=True)


class NumberedCanvas:
    """Two-pass canvas to dynamically compute and draw total page numbers and running headers."""
    @classmethod
    def create(cls):
        from reportlab.pdfgen import canvas
        from reportlab.lib import colors

        class _NumberedCanvas(canvas.Canvas):
            def __init__(self, *args, **kwargs):
                super().__init__(*args, **kwargs)
                self._saved_page_states = []

            def showPage(self):
                self._saved_page_states.append(dict(self.__dict__))
                self._startPage()

            def save(self):
                num_pages = len(self._saved_page_states)
                for state in self._saved_page_states:
                    self.__dict__.update(state)
                    self.draw_header_footer(num_pages)
                    super().showPage()
                super().save()

            def draw_header_footer(self, page_count):
                self.saveState()
                self.setFont("Helvetica", 8)
                self.setFillColor(colors.HexColor("#64748B"))
                
                # Running Header (pages 2+)
                if self._pageNumber > 1:
                    self.drawString(40, 760, "ResearchAI  |  Comprehensive Scientific Intelligence Report")
                    self.setStrokeColor(colors.HexColor("#E2E8F0"))
                    self.setLineWidth(0.5)
                    self.line(40, 754, 572, 754)

                # Running Footer (all pages)
                self.setStrokeColor(colors.HexColor("#E2E8F0"))
                self.setLineWidth(0.5)
                self.line(40, 40, 572, 40)
                
                self.drawString(40, 28, "Verified Multi-Agent Synthesis  •  Autonomous Research Fleet")
                page_text = f"Page {self._pageNumber} of {page_count}"
                self.drawRightString(572, 28, page_text)
                self.restoreState()

        return _NumberedCanvas


class ReportGenerator:
    """Generates enterprise research reports in PDF, DOCX, Markdown, and HTML formats."""

    async def generate_all(
        self,
        report_id: UUID,
        content: Dict[str, Any],
        title: str,
        query: str,
        citation_style: str = "apa",
    ) -> Dict[str, Optional[str]]:
        """Generate all export formats and return file paths."""
        paths = {}
        report_dir = REPORTS_DIR / str(report_id)
        report_dir.mkdir(parents=True, exist_ok=True)

        # 1. Markdown (.md)
        md_path = report_dir / "report.md"
        try:
            md_content = self._generate_markdown(content, title, query, citation_style)
            md_path.write_text(md_content, encoding="utf-8")
            paths["markdown"] = str(md_path)
        except Exception as e:
            logger.error("Markdown generation failed", error=str(e))
            paths["markdown"] = None

        # 2. HTML (.html)
        html_path = report_dir / "report.html"
        try:
            html_content = self._generate_html(content, title, query, citation_style)
            html_path.write_text(html_content, encoding="utf-8")
            paths["html"] = str(html_path)
        except Exception as e:
            logger.error("HTML generation failed", error=str(e))
            paths["html"] = None

        # 3. PDF (.pdf) via ReportLab
        pdf_path = report_dir / "report.pdf"
        try:
            self._generate_pdf_reportlab(content, title, query, citation_style, str(pdf_path))
            paths["pdf"] = str(pdf_path)
        except Exception as e:
            logger.error("PDF generation failed", error=str(e))
            paths["pdf"] = None

        # 4. Word Document (.docx)
        docx_path = report_dir / "report.docx"
        try:
            self._generate_docx(content, title, query, citation_style, str(docx_path))
            paths["docx"] = str(docx_path)
        except Exception as e:
            logger.error("DOCX generation failed", error=str(e))
            paths["docx"] = None

        return paths

    def generate_single_format(
        self,
        report_id: UUID,
        format_name: str,
        content: Dict[str, Any],
        title: str,
        query: str,
        citation_style: str = "apa",
    ) -> Optional[str]:
        """Generate a single format on-demand."""
        report_dir = REPORTS_DIR / str(report_id)
        report_dir.mkdir(parents=True, exist_ok=True)

        if format_name == "markdown":
            p = report_dir / "report.md"
            p.write_text(self._generate_markdown(content, title, query, citation_style), encoding="utf-8")
            return str(p)
        elif format_name == "html":
            p = report_dir / "report.html"
            p.write_text(self._generate_html(content, title, query, citation_style), encoding="utf-8")
            return str(p)
        elif format_name == "pdf":
            p = report_dir / "report.pdf"
            self._generate_pdf_reportlab(content, title, query, citation_style, str(p))
            return str(p)
        elif format_name == "docx":
            p = report_dir / "report.docx"
            self._generate_docx(content, title, query, citation_style, str(p))
            return str(p)
        return None

    def _generate_markdown(
        self, content: Dict, title: str, query: str, citation_style: str
    ) -> str:
        references = content.get("references", [])
        src_count = content.get("source_count", len(references))

        md = f"# {title}\n\n"
        md += f"**Research Inquiry:** {query}\n\n"
        md += f"**Date of Synthesis:** {datetime.now().strftime('%Y-%m-%d %H:%M UTC')}\n\n"
        md += f"**Evidence Synthesized:** {src_count} verified literature records\n\n"
        md += "---\n\n"

        if content.get("executive_summary"):
            md += "## Executive Summary\n\n"
            md += f"{content.get('executive_summary', '')}\n\n"

        if content.get("background_and_context"):
            md += "## Background & Foundational Context\n\n"
            md += f"{content.get('background_and_context', '')}\n\n"

        if content.get("methodology"):
            md += "## Methodology & Verification Protocol\n\n"
            md += f"{content.get('methodology', '')}\n\n"

        if content.get("findings"):
            md += "## Comprehensive Empirical Findings\n\n"
            for i, finding in enumerate(content.get("findings", []), 1):
                md += f"### {i}. {finding.get('section', '')}\n\n"
                md += f"{finding.get('content', '')}\n\n"
                if finding.get("key_takeaways"):
                    md += "**Key Empirical Takeaways:**\n"
                    for t in finding.get("key_takeaways", []):
                        md += f"- {t}\n"
                    md += "\n"
                if finding.get("evidence"):
                    ev_list = finding.get("evidence", [])
                    ev_str = ", ".join(ev_list) if isinstance(ev_list, list) else str(ev_list)
                    md += f"**Supporting Evidence:** {ev_str}\n\n"

        if content.get("analysis"):
            md += "## In-Depth Thematic Analysis\n\n"
            md += f"{content.get('analysis', '')}\n\n"

        if content.get("comparisons"):
            md += "## Evidence Matrix & Comparative Dimensions\n\n"
            for comp in content.get("comparisons", []):
                md += f"### {comp.get('aspect', '')}\n\n"
                if comp.get("analysis"):
                    md += f"{comp.get('analysis')}\n\n"
                for pos in comp.get("positions", []):
                    srcs = f" ({', '.join(pos.get('sources', []))})" if pos.get("sources") else ""
                    ev = f" — *{pos.get('evidence')}*" if pos.get("evidence") else ""
                    md += f"- **{pos.get('stance')}**{srcs}{ev}\n"
                md += "\n"

        if content.get("practical_implications"):
            md += "## Practical, Clinical & Industrial Implications\n\n"
            md += f"{content.get('practical_implications', '')}\n\n"

        if content.get("conclusions"):
            md += "## Conclusions & Consensus Milestones\n\n"
            md += f"{content.get('conclusions', '')}\n\n"

        if content.get("limitations"):
            md += "## Limitations & Methodological Constraints\n\n"
            md += f"{content.get('limitations', '')}\n\n"

        if content.get("future_directions"):
            md += "## Future Research Directions & Strategic Roadmap\n\n"
            fd = content.get("future_directions")
            if isinstance(fd, list):
                for item in fd:
                    md += f"- {item}\n"
                md += "\n"
            else:
                md += f"{fd}\n\n"

        if references:
            md += "## References & Bibliography\n\n"
            for i, ref in enumerate(references, 1):
                citation = ref.get(f"citation_{citation_style}") or ref.get("apa") or ref.get("title", "")
                md += f"{i}. {citation}\n"

        return md

    def _generate_html(
        self, content: Dict, title: str, query: str, citation_style: str
    ) -> str:
        references = content.get("references", [])
        src_count = content.get("source_count", len(references))

        findings_html = ""
        for i, finding in enumerate(content.get("findings", []), 1):
            takeaways_html = ""
            if finding.get("key_takeaways"):
                takeaways_html = "<div class='takeaways'><strong>Key Empirical Takeaways:</strong><ul>" + "".join(f"<li>{t}</li>" for t in finding.get("key_takeaways", [])) + "</ul></div>"
            ev_list = finding.get("evidence", [])
            ev_str = ", ".join(ev_list) if isinstance(ev_list, list) else str(ev_list)
            evidence_html = f"<p class='evidence'><em>Supporting Evidence:</em> {ev_str}</p>" if ev_str else ""
            findings_html += f"<h3>{i}. {finding.get('section', '')}</h3><p>{finding.get('content', '')}</p>{takeaways_html}{evidence_html}"

        comparisons_html = ""
        for comp in content.get("comparisons", []):
            analysis_p = f"<p><em>{comp.get('analysis')}</em></p>" if comp.get("analysis") else ""
            comparisons_html += f"<h3>{comp.get('aspect', '')}</h3>{analysis_p}<ul>"
            for pos in comp.get("positions", []):
                srcs = f" ({', '.join(pos.get('sources', []))})" if pos.get("sources") else ""
                ev_str = f" — <em>{pos.get('evidence')}</em>" if pos.get("evidence") else ""
                comparisons_html += f"<li><strong>{pos.get('stance')}</strong>{srcs}{ev_str}</li>"
            comparisons_html += "</ul>"

        refs_html = "".join(
            f"<li>{ref.get(f'citation_{citation_style}') or ref.get('apa') or ref.get('title', '')}</li>"
            for ref in references
        )

        bg_html = f"<h2>Background & Foundational Context</h2><p>{content.get('background_and_context', '')}</p>" if content.get("background_and_context") else ""
        method_html = f"<h2>Methodology & Verification Protocol</h2><p>{content.get('methodology', '')}</p>" if content.get("methodology") else ""
        implications_html = f"<h2>Practical, Clinical & Industrial Implications</h2><p>{content.get('practical_implications', '')}</p>" if content.get("practical_implications") else ""
        conclusions_html = f"<h2>Conclusions & Consensus Milestones</h2><p>{content.get('conclusions', '')}</p>" if content.get("conclusions") else ""
        limitations_html = f"<h2>Limitations & Methodological Constraints</h2><p>{content.get('limitations', '')}</p>" if content.get("limitations") else ""
        
        future_data = content.get("future_directions")
        future_body = ""
        if isinstance(future_data, list):
            future_body = "<ul>" + "".join(f"<li>{f}</li>" for f in future_data) + "</ul>"
        elif future_data:
            future_body = f"<p>{future_data}</p>"
        future_html = f"<h2>Future Research Directions & Strategic Roadmap</h2>{future_body}" if future_body else ""

        return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title}</title>
<style>
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif; line-height: 1.8; color: #1e293b; max-width: 960px; margin: 0 auto; padding: 40px 24px; background: #fafafa; }}
  .container {{ background: #ffffff; padding: 48px; border-radius: 16px; border: 1px solid #e2e8f0; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05); }}
  .badge {{ font-size: 0.75rem; font-weight: 700; color: #059669; text-transform: uppercase; letter-spacing: 0.05em; margin-bottom: 0.5rem; }}
  h1 {{ font-size: 1.8rem; color: #0f172a; margin-bottom: 0.5rem; border-bottom: 3px solid #059669; padding-bottom: 1rem; }}
  h2 {{ font-size: 1.3rem; color: #065f46; margin: 2rem 0 1rem; border-bottom: 1px solid #e2e8f0; padding-bottom: 0.4rem; }}
  h3 {{ font-size: 1.1rem; color: #334155; margin: 1.4rem 0 0.5rem; }}
  p {{ margin-bottom: 1rem; text-align: justify; color: #334155; }}
  .meta {{ color: #64748b; font-size: 0.85rem; margin-bottom: 2rem; background: #f8fafc; padding: 1rem; border-radius: 8px; border: 1px solid #e2e8f0; }}
  .executive-summary {{ margin-bottom: 2rem; padding: 1.5rem; background: #f0fdf4; border-radius: 8px; border-left: 4px solid #059669; }}
  .takeaways {{ background: #f0fdf4; border: 1px solid #a7f3d0; padding: 1rem; border-radius: 8px; margin: 0.75rem 0; font-size: 0.9rem; }}
  .evidence {{ font-size: 0.85rem; color: #64748b; margin-top: 0.5rem; }}
  ul {{ padding-left: 1.5rem; margin-bottom: 1rem; }}
  li {{ margin-bottom: 0.5rem; color: #334155; }}
  .references {{ font-size: 0.9rem; margin-top: 2rem; }}
  .references ol {{ counter-reset: ref-counter; list-style: none; padding: 0; }}
  .references ol li {{ counter-increment: ref-counter; padding-left: 2.2rem; position: relative; margin-bottom: 0.75rem; }}
  .references ol li::before {{ content: "[" counter(ref-counter) "]"; position: absolute; left: 0; color: #059669; font-weight: bold; }}
  @media print {{ body {{ max-width: 100%; background: #ffffff; padding: 0; }} .container {{ border: none; box-shadow: none; padding: 0; }} }}
</style>
</head>
<body>
<div class="container">
<div class="badge">Comprehensive Scientific Intelligence Report</div>
<h1>{title}</h1>
<div class="meta">
  <strong>Research Inquiry:</strong> {query}<br>
  <strong>Date of Synthesis:</strong> {datetime.now().strftime('%B %d, %Y at %H:%M UTC')}<br>
  <strong>Evidence Synthesized:</strong> {src_count} verified literature records
</div>

<div class="executive-summary">
<h2>Executive Summary</h2>
<p>{content.get('executive_summary', '')}</p>
</div>

{bg_html}

{method_html}

<h2>Comprehensive Empirical Findings</h2>
{findings_html}

<h2>In-Depth Thematic Analysis</h2>
<p>{content.get('analysis', '')}</p>

{'<h2>Evidence Matrix & Comparative Dimensions</h2>' + comparisons_html if comparisons_html else ''}

{implications_html}

{conclusions_html}

{limitations_html}

{future_html}

<div class="references">
<h2>References & Scholarly Bibliography</h2>
<ol>{refs_html}</ol>
</div>
</div>
</body>
</html>"""

    def _generate_pdf_reportlab(
        self, content: Dict, title: str, query: str, citation_style: str, output_path: str
    ) -> None:
        """High-density, professional ReportLab PDF generator matching frontend preview layout."""
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_RIGHT

        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            leftMargin=40,
            rightMargin=40,
            topMargin=46,
            bottomMargin=46,
        )

        styles = getSampleStyleSheet()

        # Palette
        c_primary = colors.HexColor("#065F46")      # Deep Emerald
        c_secondary = colors.HexColor("#059669")    # Accent Emerald
        c_dark = colors.HexColor("#0F172A")         # Heading Slate
        c_body = colors.HexColor("#334155")         # Body Slate
        c_muted = colors.HexColor("#64748B")        # Muted Slate
        c_box_bg = colors.HexColor("#F0FDF4")       # Light Emerald
        c_box_border = colors.HexColor("#A7F3D0")   # Border Emerald
        c_gray_bg = colors.HexColor("#F8FAFC")      # Light Gray
        c_gray_border = colors.HexColor("#E2E8F0")

        # Typography Styles
        title_style = ParagraphStyle(
            "DocTitle",
            parent=styles["Title"],
            fontSize=15,
            leading=20,
            textColor=c_dark,
            alignment=TA_LEFT,
            fontName="Helvetica-Bold",
            spaceAfter=6,
        )

        badge_style = ParagraphStyle(
            "BadgeText",
            parent=styles["Normal"],
            fontSize=7.5,
            leading=9.5,
            textColor=c_secondary,
            fontName="Helvetica-Bold",
            spaceAfter=3,
        )

        h1_style = ParagraphStyle(
            "H1_Custom",
            parent=styles["Heading1"],
            fontSize=11.5,
            leading=15,
            textColor=c_primary,
            fontName="Helvetica-Bold",
            spaceBefore=12,
            spaceAfter=5,
            keepWithNext=True,
        )

        h2_style = ParagraphStyle(
            "H2_Custom",
            parent=styles["Heading2"],
            fontSize=10,
            leading=13.5,
            textColor=c_dark,
            fontName="Helvetica-Bold",
            spaceBefore=9,
            spaceAfter=4,
            keepWithNext=True,
        )

        body_style = ParagraphStyle(
            "Body_Custom",
            parent=styles["BodyText"],
            fontSize=9,
            leading=13.5,
            textColor=c_body,
            alignment=TA_JUSTIFY,
            spaceAfter=6,
            fontName="Helvetica",
        )

        callout_body = ParagraphStyle(
            "CalloutBody",
            parent=body_style,
            fontSize=9,
            leading=13.5,
            textColor=colors.HexColor("#064E3B"),
            spaceAfter=4,
        )

        bullet_style = ParagraphStyle(
            "Bullet_Custom",
            parent=styles["Normal"],
            fontSize=8.5,
            leading=12.5,
            textColor=c_body,
            leftIndent=12,
            spaceAfter=3,
            fontName="Helvetica",
        )

        meta_label = ParagraphStyle(
            "MetaLabel",
            parent=styles["Normal"],
            fontSize=8,
            leading=11.5,
            textColor=c_muted,
            fontName="Helvetica",
        )

        story = []

        # 1. Header Badge & Title
        story.append(Paragraph("COMPREHENSIVE SCIENTIFIC INTELLIGENCE REPORT", badge_style))
        story.append(Paragraph(title, title_style))
        story.append(HRFlowable(width="100%", thickness=1.5, color=c_secondary, spaceBefore=2, spaceAfter=6))

        # 2. Metadata Card (Compact table that fits easily on page 1)
        references = content.get("references", [])
        src_count = content.get("source_count", len(references))
        meta_left = f"<b>Research Inquiry:</b> {query}"
        meta_right = f"<b>Date:</b> {datetime.now().strftime('%B %d, %Y')} | <b>Sources Synthesized:</b> {src_count} records"
        
        meta_tbl = Table(
            [[Paragraph(meta_left, meta_label), Paragraph(meta_right, ParagraphStyle("MetaR", parent=meta_label, alignment=TA_RIGHT))]],
            colWidths=[230, 302]
        )
        meta_tbl.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), c_gray_bg),
            ('BOX', (0, 0), (-1, -1), 0.75, c_gray_border),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('PADDING', (0, 0), (-1, -1), 5),
        ]))
        story.append(meta_tbl)
        story.append(Spacer(1, 8))

        # 3. Executive Summary (Flowing paragraphs with light background)
        exec_summary = content.get("executive_summary", "")
        if exec_summary:
            story.append(Paragraph("Executive Summary", h1_style))
            for p_text in exec_summary.split("\n\n"):
                if p_text.strip():
                    p_flow = Paragraph(p_text.strip(), callout_body)
                    p_tbl = Table([[p_flow]], colWidths=[532])
                    p_tbl.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, -1), c_box_bg),
                        ('LINELEFT', (0, 0), (-1, -1), 3, c_secondary),
                        ('PADDING', (0, 0), (-1, -1), 5),
                    ]))
                    story.append(p_tbl)
                    story.append(Spacer(1, 3))
            story.append(Spacer(1, 4))

        # Helper for multi-paragraph or list sections
        def render_section(heading, data):
            if not data:
                return
            story.append(Paragraph(heading, h1_style))
            if isinstance(data, str):
                for p_text in data.split("\n\n"):
                    if p_text.strip():
                        story.append(Paragraph(p_text.strip(), body_style))
            elif isinstance(data, list):
                for item in data:
                    if isinstance(item, str) and item.strip():
                        story.append(Paragraph(f"• {item.strip()}", bullet_style))
                    elif isinstance(item, dict):
                        story.append(Paragraph(f"• <b>{item.get('title', '')}</b>: {item.get('description', item.get('content', ''))}", bullet_style))
            story.append(Spacer(1, 4))

        # 4. Background & Foundational Context
        render_section("Background & Foundational Context", content.get("background_and_context"))

        # 5. Methodology & Verification Protocol
        render_section("Methodology & Verification Protocol", content.get("methodology"))

        # 6. Comprehensive Empirical Findings
        findings = content.get("findings", [])
        if findings:
            story.append(Paragraph("Comprehensive Empirical Findings", h1_style))
            for i, f in enumerate(findings, 1):
                sec_title = f"{i}. {f.get('section', '')}"
                story.append(Paragraph(sec_title, h2_style))
                
                # Content paragraphs
                f_content = f.get("content", "")
                if isinstance(f_content, str):
                    for p_text in f_content.split("\n\n"):
                        if p_text.strip():
                            story.append(Paragraph(p_text.strip(), body_style))
                
                # Takeaways
                takeaways = f.get("key_takeaways", [])
                if takeaways:
                    t_elements = [Paragraph("<b>Key Empirical Takeaways:</b>", ParagraphStyle("TKW", parent=bullet_style, textColor=c_primary, spaceAfter=2))]
                    for t in takeaways:
                        t_elements.append(Paragraph(f"• {t}", bullet_style))
                    t_tbl = Table([[t_elements]], colWidths=[532])
                    t_tbl.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F0FDF4")),
                        ('BOX', (0, 0), (-1, -1), 0.5, c_box_border),
                        ('PADDING', (0, 0), (-1, -1), 5),
                    ]))
                    story.append(t_tbl)
                    story.append(Spacer(1, 3))

                # Supporting Evidence
                if f.get("evidence"):
                    ev_list = f.get("evidence", [])
                    ev_text = f"<b>Supporting Evidence:</b> {', '.join(ev_list) if isinstance(ev_list, list) else ev_list}"
                    story.append(Paragraph(ev_text, meta_label))
                story.append(Spacer(1, 6))

        # 7. In-Depth Thematic Analysis
        render_section("In-Depth Thematic Analysis", content.get("analysis"))

        # 8. Evidence Matrix & Comparative Dimensions
        comparisons = content.get("comparisons", [])
        if comparisons:
            story.append(Paragraph("Evidence Matrix & Comparative Dimensions", h1_style))
            for c in comparisons:
                story.append(Paragraph(c.get("aspect", ""), h2_style))
                if c.get("analysis"):
                    story.append(Paragraph(f"<i>{c.get('analysis')}</i>", body_style))
                for pos in c.get("positions", []):
                    ev = f" — <i>{pos.get('evidence')}</i>" if pos.get("evidence") else ""
                    srcs = f" ({', '.join(pos.get('sources', []))})" if pos.get("sources") else ""
                    story.append(Paragraph(f"• <b>{pos.get('stance')}</b>{srcs}{ev}", bullet_style))
                story.append(Spacer(1, 5))

        # 9. Practical, Clinical & Industrial Implications
        render_section("Practical, Clinical & Industrial Implications", content.get("practical_implications"))

        # 10. Conclusions & Consensus Milestones
        render_section("Conclusions & Consensus Milestones", content.get("conclusions"))

        # 11. Limitations & Methodological Constraints
        render_section("Limitations & Methodological Constraints", content.get("limitations"))

        # 12. Future Research Directions & Strategic Roadmap
        render_section("Future Research Directions & Strategic Roadmap", content.get("future_directions"))

        # 13. References & Scholarly Bibliography
        if references:
            story.append(Paragraph("References & Scholarly Bibliography", h1_style))
            for i, ref in enumerate(references, 1):
                citation = ref.get(f"citation_{citation_style}") or ref.get("apa") or ref.get("title", "")
                story.append(Paragraph(f"[{i}] {citation}", bullet_style))

        # Build using NumberedCanvas
        doc.build(story, canvasmaker=NumberedCanvas.create())

    def _generate_docx(
        self, content: Dict, title: str, query: str, citation_style: str, output_path: str
    ) -> None:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Metadata
        meta = doc.add_paragraph()
        meta.add_run("Research Inquiry: ").bold = True
        meta.add_run(query)
        meta.add_run(f"\nGenerated: {datetime.now().strftime('%B %d, %Y')}")
        meta.add_run(f"\nPrimary Sources Synthesized: {content.get('source_count', len(content.get('references', [])))}")

        doc.add_page_break()

        # Sections
        sections = [
            ("Executive Summary", content.get("executive_summary", "")),
            ("Background & Context", content.get("background_and_context", "")),
            ("Methodology & Verification Protocol", content.get("methodology", "")),
            ("In-Depth Thematic Analysis", content.get("analysis", "")),
            ("Practical & Translational Implications", content.get("practical_implications", "")),
            ("Conclusions & Strategic Roadmap", content.get("conclusions", "")),
            ("Limitations & Constraints", content.get("limitations", "")),
        ]

        for section_title, section_content in sections:
            if section_content:
                doc.add_heading(section_title, level=1)
                if isinstance(section_content, str):
                    doc.add_paragraph(section_content)
                elif isinstance(section_content, list):
                    for item in section_content:
                        doc.add_paragraph(f"• {item}", style="List Bullet")

        # Future Directions
        fd = content.get("future_directions")
        if fd:
            doc.add_heading("Future Research Directions & Strategic Roadmap", level=1)
            if isinstance(fd, list):
                for item in fd:
                    doc.add_paragraph(f"• {item}", style="List Bullet")
            else:
                doc.add_paragraph(str(fd))

        # Findings
        if content.get("findings"):
            doc.add_heading("Findings & Empirical Insights", level=1)
            for i, finding in enumerate(content.get("findings", []), 1):
                doc.add_heading(f"{i}. {finding.get('section', '')}", level=2)
                doc.add_paragraph(finding.get("content", ""))
                if finding.get("key_takeaways"):
                    for t in finding.get("key_takeaways", []):
                        doc.add_paragraph(f"• {t}", style="List Bullet")

        # References
        references = content.get("references", [])
        if references:
            doc.add_heading("References", level=1)
            for i, ref in enumerate(references, 1):
                citation = ref.get(f"citation_{citation_style}") or ref.get("apa") or ref.get("title", "")
                doc.add_paragraph(f"[{i}] {citation}", style="List Number")

        doc.save(output_path)


report_generator = ReportGenerator()
